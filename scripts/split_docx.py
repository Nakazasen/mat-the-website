import sys
import re
from pathlib import Path
from docx import Document

CHAPTER_HEADER_RE = re.compile(
    r"^(?:.*?)Ch[uư][oơ]ng\s+(\d+)\s*[:\-–—]?\s*(.*)$",
    re.IGNORECASE,
)

def copy_paragraph(original_paragraph, new_document):
    """Kopieren eines Absatzes mit grundlegender Formatierung."""
    new_p = new_document.add_paragraph()
    new_p.style = original_paragraph.style
    new_p.alignment = original_paragraph.alignment
    
    for run in original_paragraph.runs:
        new_run = new_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
        if run.font.size:
            new_run.font.size = run.font.size
        # Not copying all font properties to keep it simple, but this should cover 99%

def main():
    input_path = Path(r"C:\ProgramData\Sandbox\Web_matthesinhhoanguyco\mat-the-website\scripts\mat-the-truyen\File gốc 1-767.docx")
    out_dir = input_path.parent
    
    print(f"Loading document: {input_path}")
    if not input_path.exists():
        print(f"Error: {input_path} does not exist!")
        sys.exit(1)
        
    doc = Document(str(input_path))
    print(f"Loaded {len(doc.paragraphs)} paragraphs.")
    
    current_number = None
    current_doc = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = CHAPTER_HEADER_RE.match(text)
        
        # Determine if it's a chapter header and short enough to not be regular text
        if match and len(text) < 150:
            # We found a new chapter
            if current_doc is not None and current_number is not None:
                # Save previous chapter
                out_name = f"Chương {current_number}.docx"
                out_path = out_dir / out_name
                print(f"Saving {out_name}...")
                current_doc.save(str(out_path))
            
            # Start new document
            current_number = int(match.group(1))
            current_doc = Document()
            print(f"Found Chapter {current_number}")
            
        # Add paragraph to current document
        if current_doc is not None:
            copy_paragraph(para, current_doc)
            
    # Save the last document
    if current_doc is not None and current_number is not None:
        out_name = f"Chương {current_number}.docx"
        out_path = out_dir / out_name
        print(f"Saving last chapter: {out_name}...")
        current_doc.save(str(out_path))
        
    print("All chapters extracted successfully.")

if __name__ == "__main__":
    main()
